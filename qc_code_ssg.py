# -*- coding: utf-8 -*-
"""
QC Code SSG — Vertex AI Gemini 2.5
OLD FORMAT RESTORED + EDITORIAL SAFETY FIXES
"""
FACT_CACHE = {}

# ===================== CORE =====================
import re
import os
import json
import base64
import requests
import hashlib
import tempfile
import sqlite3
import uuid
import io
import streamlit as st
IMPORT_ONLY = os.environ.get("QC_SSG_IMPORT_ONLY") == "1"
try:
    import extra_streamlit_components as stx
except Exception:
    if IMPORT_ONLY:
        class _DummyCookieManager:
            def get(self, *args, **kwargs):
                return ""

            def set(self, *args, **kwargs):
                return None

            def delete(self, *args, **kwargs):
                return None

        class _DummySTX:
            CookieManager = _DummyCookieManager

        stx = _DummySTX()
    else:
        raise
import html
import time
from datetime import datetime, timezone, timedelta
from concurrent.futures import ThreadPoolExecutor, as_completed
from bs4 import BeautifulSoup
from google.oauth2 import service_account
from google.auth.transport.requests import Request as GoogleAuthRequest
from difflib import SequenceMatcher
from urllib.parse import quote


# ===================== NLP =====================
import spacy

# ===================== SPELL & FREQUENCY =====================
from spellchecker import SpellChecker
from wordfreq import zipf_frequency
from symspellpy import SymSpell, Verbosity
from importlib.resources import files as pkg_files

# ===================== DOCX =====================
from docx import Document

# ===================== GRAMMAR =====================
import language_tool_python


# ===================== GEN AI CLIENT =====================
from google import genai
from google.genai import types as genai_types


# =================================================
# STREAMLIT CONFIG
# =================================================
if not IMPORT_ONLY:
    st.set_page_config(page_title="Article QC Tool (Gemini 2.5)", layout="wide")

def _secret(name: str, default=""):
    try:
        return st.secrets[name]
    except Exception:
        return default

ALLOWED_EMAIL_DOMAIN = str(_secret("ALLOWED_EMAIL_DOMAIN", "jagrannewmedia.com")).strip().lower()
APP_ACCESS_SUPPORT_TEXT = str(
    _secret(
        "APP_ACCESS_SUPPORT_TEXT",
        "Enter your Jagran New Media email address to continue.",
    )
).strip()
ADMIN_EMAIL = "kartikay.khosla@jagrannewmedia.com"
HISTORY_DB_PATH = os.path.join(os.path.dirname(__file__), ".app_history.sqlite3")
HISTORY_SPREADSHEET_ID = str(_secret("HISTORY_SPREADSHEET_ID", "")).strip()
SESSION_QUERY_KEY = "_jnm_session"
SESSION_COOKIE_KEY = "_jnm_session"
SESSION_EMAIL_COOKIE_KEY = "_jnm_email"
SESSION_EXP_COOKIE_KEY = "_jnm_session_exp"
SESSION_TTL_HOURS = 24
SESSION_REFRESH_WINDOW_MINUTES = 15

HISTORY_HEADERS = {
    "login_events": ["ts_utc", "app", "email"],
    "analysis_runs": [
        "run_id",
        "ts_utc",
        "app",
        "email",
        "source_type",
        "source_identity",
        "source_label",
        "analysis_key",
        "iteration",
        "spelling_count",
        "grammar_count",
        "editorial_count",
        "fact_count",
        "total_count",
    ],
    "access_sessions": [
        "ts_utc",
        "app",
        "token_hash",
        "email",
        "event_type",
        "expires_ts_utc",
        "last_seen_ts_utc",
    ],
}

def _email_allowed(email: str) -> bool:
    return (email or "").strip().lower().endswith(f"@{ALLOWED_EMAIL_DOMAIN}")

def _normalise_username(username: str) -> str:
    value = (username or "").strip().lower().replace(" ", "")
    if "@" in value:
        return ""
    return value

def _build_email_from_username(username: str) -> str:
    username = _normalise_username(username)
    if not username:
        return ""
    return f"{username}@{ALLOWED_EMAIL_DOMAIN}"

def _email_access_granted() -> bool:
    email = st.session_state.get("_email_access_email", "")
    return bool(st.session_state.get("_email_access_granted")) and _email_allowed(email)

def _clear_email_access():
    st.session_state.pop("_email_access_granted", None)
    st.session_state.pop("_email_access_email", None)
    _clear_pending_analysis_state()

def _current_access_email() -> str:
    return (st.session_state.get("_email_access_email") or "").strip().lower()

def _is_admin_user() -> bool:
    return _current_access_email() == ADMIN_EMAIL

def _history_conn():
    conn = sqlite3.connect(HISTORY_DB_PATH)
    conn.row_factory = sqlite3.Row
    return conn

def _history_uses_sheets() -> bool:
    return bool(HISTORY_SPREADSHEET_ID)

def _sheet_tab_title(app_name: str, kind: str) -> str:
    safe_app = re.sub(r"[^A-Za-z0-9_\\-]", "_", app_name or "app").strip("_") or "app"
    safe_kind = re.sub(r"[^A-Za-z0-9_\\-]", "_", kind or "history").strip("_") or "history"
    return f"{safe_app}_{safe_kind}"[:95]

def _history_headers(kind: str):
    return HISTORY_HEADERS.get(kind, [])

def _parse_utc_iso(value: str):
    text = (value or "").strip()
    if not text:
        return None
    if text.endswith("Z"):
        text = text[:-1] + "+00:00"
    try:
        parsed = datetime.fromisoformat(text)
        if parsed.tzinfo is None:
            parsed = parsed.replace(tzinfo=timezone.utc)
        return parsed.astimezone(timezone.utc)
    except Exception:
        return None

def _safe_int(value, default=0):
    try:
        return int(value)
    except Exception:
        return default

def _utc_now():
    return datetime.now(timezone.utc)

def _hash_session_token(token: str) -> str:
    return hashlib.sha256(f"{ALLOWED_EMAIL_DOMAIN}:{token}".encode("utf-8")).hexdigest()

def _get_cookie_manager():
    manager = st.session_state.get("_cookie_manager")
    if manager is None:
        manager = stx.CookieManager()
        st.session_state["_cookie_manager"] = manager
    return manager

def _get_context_cookie(name: str) -> str:
    try:
        cookies = getattr(st.context, "cookies", None)
        if cookies is None:
            return ""
        value = cookies.get(name, "")
        return (value or "").strip()
    except Exception:
        return ""

def _get_cookie_value(name: str) -> str:
    value = _get_context_cookie(name)
    if value:
        return value
    try:
        value = _get_cookie_manager().get(name)
        return (value or "").strip()
    except Exception:
        return ""

def _set_cookie_value(name: str, value: str):
    try:
        _get_cookie_manager().set(
            name,
            value,
            expires_at=datetime.now() + timedelta(hours=SESSION_TTL_HOURS),
            key=f"set-cookie-{name}",
        )
    except Exception:
        pass

def _clear_cookie_value(name: str):
    try:
        _get_cookie_manager().delete(name, key=f"delete-cookie-{name}")
    except Exception:
        pass

def _get_session_query_token() -> str:
    try:
        value = st.query_params.get(SESSION_QUERY_KEY, "")
        if isinstance(value, list):
            return (value[0] or "").strip()
        return (value or "").strip()
    except Exception:
        return ""

def _set_session_query_token(token: str):
    try:
        st.query_params[SESSION_QUERY_KEY] = token
    except Exception:
        pass

def _clear_session_query_token():
    try:
        st.query_params.pop(SESSION_QUERY_KEY, None)
    except Exception:
        pass

def _get_session_cookie_token() -> str:
    return _get_cookie_value(SESSION_COOKIE_KEY)

def _set_session_cookie_token(token: str):
    _set_cookie_value(SESSION_COOKIE_KEY, token)

def _clear_session_cookie_token():
    _clear_cookie_value(SESSION_COOKIE_KEY)

def _get_session_identity_cookie_email() -> str:
    return _get_cookie_value(SESSION_EMAIL_COOKIE_KEY).lower()

def _get_session_identity_cookie_expiry() -> str:
    return _get_cookie_value(SESSION_EXP_COOKIE_KEY)

def _set_session_identity_cookies(email: str, expires_iso: str):
    _set_cookie_value(SESSION_EMAIL_COOKIE_KEY, (email or "").strip().lower())
    _set_cookie_value(SESSION_EXP_COOKIE_KEY, expires_iso)

def _clear_session_identity_cookies():
    _clear_cookie_value(SESSION_EMAIL_COOKIE_KEY)
    _clear_cookie_value(SESSION_EXP_COOKIE_KEY)

def _restore_identity_cookie_session() -> bool:
    email = _get_session_identity_cookie_email()
    expires_at = _parse_utc_iso(_get_session_identity_cookie_expiry())
    if not email or not _email_allowed(email) or not expires_at or expires_at <= _utc_now():
        _clear_session_identity_cookies()
        return False
    refreshed_expiry = (_utc_now() + timedelta(hours=SESSION_TTL_HOURS)).isoformat()
    _set_session_identity_cookies(email, refreshed_expiry)
    st.session_state["_email_access_granted"] = True
    st.session_state["_email_access_email"] = email
    return True

def _sqlite_rows(query: str, params=()):
    try:
        ensure_history_db()
        with _history_conn() as conn:
            return [dict(row) for row in conn.execute(query, params).fetchall()]
    except Exception:
        return []

def _load_service_account_info():
    if "GCP_SERVICE_ACCOUNT_JSON_B64" not in st.secrets:
        st.error("❌ GCP_SERVICE_ACCOUNT_JSON_B64 not set in Streamlit secrets")
        st.stop()
    decoded = base64.b64decode(
        st.secrets["GCP_SERVICE_ACCOUNT_JSON_B64"]
    ).decode("utf-8")
    return json.loads(decoded)

def _get_scoped_service_account_credentials(scopes):
    creds_dict = _load_service_account_info()
    with open(CRED_PATH, "w") as f:
        json.dump(creds_dict, f)
    os.environ["GOOGLE_APPLICATION_CREDENTIALS"] = CRED_PATH
    creds = service_account.Credentials.from_service_account_info(
        creds_dict,
        scopes=list(scopes),
    )
    project_id = PROJECT_ID or str(creds_dict.get("project_id", "")).strip()
    if not project_id:
        st.error("❌ Could not determine Vertex project ID from secrets or service account JSON")
        st.stop()
    return creds, project_id, creds_dict

def _sheets_api_request(method: str, path: str = "", params=None, json_body=None):
    creds, _, _ = _get_scoped_service_account_credentials([CLOUD_PLATFORM_SCOPE, SHEETS_SCOPE])
    creds.refresh(GoogleAuthRequest())
    headers = {"Authorization": f"Bearer {creds.token}"}
    if json_body is not None:
        headers["Content-Type"] = "application/json"
    url = f"https://sheets.googleapis.com/v4/spreadsheets/{HISTORY_SPREADSHEET_ID}{path}"
    response = requests.request(
        method,
        url,
        headers=headers,
        params=params,
        json=json_body,
        timeout=30,
    )
    if response.status_code >= 400:
        raise RuntimeError(f"Google Sheets API {response.status_code}: {response.text}")
    if not response.text:
        return {}
    return response.json()

def _ensure_history_sheet(app_name: str, kind: str) -> str:
    if not _history_uses_sheets():
        return ""
    tab_title = _sheet_tab_title(app_name, kind)
    metadata = _sheets_api_request("GET", params={"fields": "sheets.properties.title"})
    titles = {
        ((sheet.get("properties") or {}).get("title") or "").strip()
        for sheet in metadata.get("sheets", [])
    }
    if tab_title not in titles:
        _sheets_api_request(
            "POST",
            ":batchUpdate",
            json_body={"requests": [{"addSheet": {"properties": {"title": tab_title}}}]},
        )
    headers = _history_headers(kind)
    encoded_range = quote(f"{tab_title}!1:1", safe="!:$")
    current = _sheets_api_request("GET", f"/values/{encoded_range}")
    current_values = current.get("values", [])
    if not current_values:
        encoded_write = quote(f"{tab_title}!A1", safe="!:$")
        _sheets_api_request(
            "PUT",
            f"/values/{encoded_write}",
            params={"valueInputOption": "RAW"},
            json_body={"range": f"{tab_title}!A1", "majorDimension": "ROWS", "values": [headers]},
        )
    return tab_title

def _sheet_read_rows(app_name: str, kind: str):
    if not _history_uses_sheets():
        return []
    try:
        tab_title = _ensure_history_sheet(app_name, kind)
        encoded_range = quote(f"{tab_title}!A:Z", safe="!:$")
        payload = _sheets_api_request("GET", f"/values/{encoded_range}")
        values = payload.get("values", [])
        headers = _history_headers(kind)
        if not values:
            return []
        start_index = 1 if values[0] == headers else 0
        rows = []
        for raw_row in values[start_index:]:
            if not any((cell or "").strip() for cell in raw_row):
                continue
            padded = list(raw_row) + [""] * max(0, len(headers) - len(raw_row))
            rows.append({header: padded[idx] if idx < len(padded) else "" for idx, header in enumerate(headers)})
        return rows
    except Exception:
        return []

def _sheet_append_row(app_name: str, kind: str, row_dict: dict) -> bool:
    if not _history_uses_sheets():
        return False
    try:
        tab_title = _ensure_history_sheet(app_name, kind)
        headers = _history_headers(kind)
        encoded_range = quote(f"{tab_title}!A:Z", safe="!:$")
        _sheets_api_request(
            "POST",
            f"/values/{encoded_range}:append",
            params={"valueInputOption": "RAW", "insertDataOption": "INSERT_ROWS"},
            json_body={"values": [[str(row_dict.get(header, "")) for header in headers]]},
        )
        return True
    except Exception:
        return False

def _sheet_login_rows(app_name: str):
    rows = _sheet_read_rows(app_name, "login_events")
    return sorted(
        [
            {"ts_utc": (row.get("ts_utc") or "").strip(), "email": (row.get("email") or "").strip().lower()}
            for row in rows
            if (row.get("email") or "").strip()
        ],
        key=lambda row: row.get("ts_utc", ""),
        reverse=True,
    )

def _sheet_analysis_rows(app_name: str):
    rows = []
    for row in _sheet_read_rows(app_name, "analysis_runs"):
        ts_value = (row.get("ts_utc") or "").strip()
        email = (row.get("email") or "").strip().lower()
        if not ts_value or not email:
            continue
        rows.append(
            {
                "run_id": (row.get("run_id") or "").strip(),
                "ts_utc": ts_value,
                "app": (row.get("app") or "").strip(),
                "email": email,
                "source_type": (row.get("source_type") or "").strip(),
                "source_identity": (row.get("source_identity") or "").strip(),
                "source_label": (row.get("source_label") or "").strip(),
                "analysis_key": (row.get("analysis_key") or "").strip(),
                "iteration": _safe_int(row.get("iteration")),
                "spelling_count": _safe_int(row.get("spelling_count")),
                "grammar_count": _safe_int(row.get("grammar_count")),
                "editorial_count": _safe_int(row.get("editorial_count")),
                "fact_count": _safe_int(row.get("fact_count")),
                "total_count": _safe_int(row.get("total_count")),
            }
        )
    return sorted(rows, key=lambda row: row.get("ts_utc", ""), reverse=True)

def _sheet_latest_session_row(app_name: str, token_hash: str):
    rows = [
        row
        for row in _sheet_read_rows(app_name, "access_sessions")
        if (row.get("token_hash") or "").strip() == token_hash
    ]
    if not rows:
        return None
    rows.sort(key=lambda row: (row.get("ts_utc") or "", row.get("last_seen_ts_utc") or ""), reverse=True)
    latest = rows[0]
    return {
        "ts_utc": (latest.get("ts_utc") or "").strip(),
        "app": (latest.get("app") or "").strip(),
        "token_hash": (latest.get("token_hash") or "").strip(),
        "email": (latest.get("email") or "").strip().lower(),
        "event_type": (latest.get("event_type") or "").strip().lower(),
        "expires_ts_utc": (latest.get("expires_ts_utc") or "").strip(),
        "last_seen_ts_utc": (latest.get("last_seen_ts_utc") or "").strip(),
    }

def _append_session_event(app_name: str, token_hash: str, email: str, event_type: str, expires_iso: str, last_seen_iso: str) -> bool:
    return _sheet_append_row(
        app_name,
        "access_sessions",
        {
            "ts_utc": _utc_now().isoformat(),
            "app": app_name,
            "token_hash": token_hash,
            "email": (email or "").strip().lower(),
            "event_type": event_type,
            "expires_ts_utc": expires_iso,
            "last_seen_ts_utc": last_seen_iso,
        },
    )

def ensure_history_db():
    try:
        with _history_conn() as conn:
            conn.execute(
                """
                CREATE TABLE IF NOT EXISTS login_events (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    ts_utc TEXT NOT NULL,
                    app TEXT NOT NULL,
                    email TEXT NOT NULL
                )
                """
            )
            conn.execute(
                """
                CREATE TABLE IF NOT EXISTS analysis_runs (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    run_id TEXT NOT NULL UNIQUE,
                    ts_utc TEXT NOT NULL,
                    app TEXT NOT NULL,
                    email TEXT NOT NULL,
                    source_type TEXT NOT NULL,
                    source_identity TEXT NOT NULL,
                    source_label TEXT NOT NULL,
                    analysis_key TEXT,
                    iteration INTEGER NOT NULL,
                    spelling_count INTEGER NOT NULL DEFAULT 0,
                    grammar_count INTEGER NOT NULL DEFAULT 0,
                    editorial_count INTEGER NOT NULL DEFAULT 0,
                    fact_count INTEGER NOT NULL DEFAULT 0,
                    total_count INTEGER NOT NULL DEFAULT 0
                )
                """
            )
            conn.execute(
                """
                CREATE TABLE IF NOT EXISTS access_sessions (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    token_hash TEXT NOT NULL UNIQUE,
                    app TEXT NOT NULL,
                    email TEXT NOT NULL,
                    created_ts_utc TEXT NOT NULL,
                    expires_ts_utc TEXT NOT NULL,
                    revoked INTEGER NOT NULL DEFAULT 0,
                    last_seen_ts_utc TEXT
                )
                """
            )
            conn.execute(
                "CREATE INDEX IF NOT EXISTS idx_login_events_app_email_ts ON login_events(app, email, ts_utc)"
            )
            conn.execute(
                "CREATE INDEX IF NOT EXISTS idx_analysis_runs_app_source_ts ON analysis_runs(app, source_identity, ts_utc)"
            )
            conn.execute(
                "CREATE INDEX IF NOT EXISTS idx_access_sessions_app_email_expiry ON access_sessions(app, email, expires_ts_utc)"
            )
    except Exception:
        pass

def _clear_pending_analysis_state():
    for key in (
        "_pending_run_id",
        "_pending_source_type",
        "_pending_source_identity",
        "_pending_source_label",
        "_pending_analysis_key",
    ):
        st.session_state.pop(key, None)

def _create_persisted_session(app_name: str, email: str):
    try:
        token = uuid.uuid4().hex + uuid.uuid4().hex
        now_iso = _utc_now().isoformat()
        expires_iso = (_utc_now() + timedelta(hours=SESSION_TTL_HOURS)).isoformat()
        token_hash = _hash_session_token(token)
        stored = False
        if _history_uses_sheets():
            stored = _append_session_event(app_name, token_hash, email, "create", expires_iso, now_iso)
        if not stored:
            ensure_history_db()
            with _history_conn() as conn:
                conn.execute(
                    """
                    INSERT INTO access_sessions (
                        token_hash, app, email, created_ts_utc, expires_ts_utc, revoked, last_seen_ts_utc
                    ) VALUES (?, ?, ?, ?, ?, 0, ?)
                    """,
                    (token_hash, app_name, (email or "").strip().lower(), now_iso, expires_iso, now_iso),
                )
        _clear_session_query_token()
        _set_session_cookie_token(token)
        _set_session_identity_cookies(email, expires_iso)
    except Exception:
        pass

def _revoke_persisted_session(app_name: str):
    token = _get_session_cookie_token() or _get_session_query_token()
    if token:
        try:
            token_hash = _hash_session_token(token)
            stored = False
            if _history_uses_sheets():
                stored = _append_session_event(
                    app_name,
                    token_hash,
                    _current_access_email(),
                    "revoke",
                    _utc_now().isoformat(),
                    _utc_now().isoformat(),
                )
            if not stored:
                ensure_history_db()
                with _history_conn() as conn:
                    conn.execute(
                        "UPDATE access_sessions SET revoked = 1 WHERE app = ? AND token_hash = ?",
                        (app_name, token_hash),
                    )
        except Exception:
            pass
    _clear_session_query_token()
    _clear_session_cookie_token()
    _clear_session_identity_cookies()

def _restore_persisted_session(app_name: str) -> bool:
    if _email_access_granted():
        return True

    token = _get_session_cookie_token() or _get_session_query_token()
    if token:
        try:
            now_iso = _utc_now().isoformat()
            token_hash = _hash_session_token(token)
            row = None
            if _history_uses_sheets():
                row = _sheet_latest_session_row(app_name, token_hash)
                if row and row.get("event_type") != "revoke":
                    expires_at = _parse_utc_iso(row.get("expires_ts_utc", ""))
                    if expires_at and expires_at > _utc_now():
                        refreshed_expiry = (_utc_now() + timedelta(hours=SESSION_TTL_HOURS)).isoformat()
                        last_seen = _parse_utc_iso(row.get("last_seen_ts_utc", ""))
                        if (not last_seen) or ((_utc_now() - last_seen) >= timedelta(minutes=SESSION_REFRESH_WINDOW_MINUTES)):
                            _append_session_event(
                                app_name,
                                token_hash,
                                row.get("email", ""),
                                "refresh",
                                refreshed_expiry,
                                now_iso,
                            )
                    else:
                        row = None
                else:
                    row = None
            else:
                ensure_history_db()
                with _history_conn() as conn:
                    sqlite_row = conn.execute(
                        """
                        SELECT email
                        FROM access_sessions
                        WHERE app = ?
                          AND token_hash = ?
                          AND revoked = 0
                          AND expires_ts_utc > ?
                        ORDER BY id DESC
                        LIMIT 1
                        """,
                        (app_name, token_hash, now_iso),
                    ).fetchone()
                    if sqlite_row:
                        refreshed_expiry = (_utc_now() + timedelta(hours=SESSION_TTL_HOURS)).isoformat()
                        conn.execute(
                            """
                            UPDATE access_sessions
                            SET last_seen_ts_utc = ?, expires_ts_utc = ?
                            WHERE app = ? AND token_hash = ?
                            """,
                            (now_iso, refreshed_expiry, app_name, token_hash),
                        )
                        row = {"email": sqlite_row["email"]}

            if row:
                email = (row.get("email") or "").strip().lower()
                if _email_allowed(email):
                    refreshed_expiry = (_utc_now() + timedelta(hours=SESSION_TTL_HOURS)).isoformat()
                    _clear_session_query_token()
                    _set_session_cookie_token(token)
                    _set_session_identity_cookies(email, refreshed_expiry)
                    st.session_state["_email_access_granted"] = True
                    st.session_state["_email_access_email"] = email
                    return True
        except Exception:
            pass

    _clear_session_query_token()
    _clear_session_cookie_token()
    return _restore_identity_cookie_session()

def queue_analysis_run(source_type: str, source_identity: str, source_label: str, analysis_key: str = ""):
    st.session_state["_pending_run_id"] = uuid.uuid4().hex
    st.session_state["_pending_source_type"] = source_type
    st.session_state["_pending_source_identity"] = source_identity
    st.session_state["_pending_source_label"] = source_label
    st.session_state["_pending_analysis_key"] = analysis_key

def _record_access_event(app_name: str, email: str):
    try:
        row = {
            "ts_utc": datetime.now(timezone.utc).isoformat(),
            "app": app_name,
            "email": (email or "").strip().lower(),
        }
        stored = _sheet_append_row(app_name, "login_events", row) if _history_uses_sheets() else False
        if not stored:
            ensure_history_db()
            with _history_conn() as conn:
                conn.execute(
                    "INSERT INTO login_events (ts_utc, app, email) VALUES (?, ?, ?)",
                    (row["ts_utc"], row["app"], row["email"]),
                )
    except Exception:
        pass

def count_markdown_rows(table_md: str, header_name: str) -> int:
    count = 0
    for line in (table_md or "").splitlines():
        row = line.strip()
        if not row.startswith("|") or row.count("|") < 2:
            continue
        parts = [part.strip() for part in row.strip("|").split("|")]
        if not parts or not any(parts):
            continue
        if all(re.fullmatch(r":?-{2,}:?", part or "") for part in parts):
            continue
        if parts[0].lower() == header_name.lower():
            continue
        count += 1
    return count

def compute_qc_score(spelling_count: int, grammar_count: int, editorial_count: int, fact_count: int) -> int:
    weighted_penalty = (
        float(spelling_count) * 0.5
        + float(grammar_count) * 1.0
        + float(editorial_count) * 0.75
        + float(fact_count) * 4.0
    )
    return max(0, min(100, int(round(100 - min(100, weighted_penalty)))))

def count_article_words(article_data) -> int:
    text = "\n".join(
        t for c, t in (article_data or []) if c in {"heading", "paragraph", "table"}
    )
    return len(re.findall(r"\S+", text))

def render_qc_score_summary(spelling_count: int, grammar_count: int, editorial_count: int, fact_count: int, has_ai_error: bool, word_count: int | None = None):
    st.markdown("### QC Summary")
    if has_ai_error:
        st.warning("QC score is unavailable because one or more AI checks failed.")
        return
    total_count = spelling_count + grammar_count + editorial_count + fact_count
    score = compute_qc_score(spelling_count, grammar_count, editorial_count, fact_count)
    score_col, word_col, spelling_col, grammar_col, editorial_col, fact_col, total_col = st.columns(7)
    score_col.metric("QC Score", f"{score}/100")
    word_col.metric("Words", word_count if word_count is not None else 0)
    spelling_col.metric("Spelling", spelling_count)
    grammar_col.metric("Grammar", grammar_count)
    editorial_col.metric("Editorial", editorial_count)
    fact_col.metric("Fact", fact_count)
    total_col.metric("Total Issues", total_count)
    st.caption("QC score is a weighted indicator based on issue counts. Fact issues carry the highest penalty.")

def log_analysis_run(app_name: str, email: str, source_type: str, source_identity: str, source_label: str,
                     analysis_key: str, spelling_count: int, grammar_count: int,
                     editorial_count: int, fact_count: int):
    run_id = st.session_state.get("_pending_run_id")
    if not run_id:
        return

    try:
        ts_utc = datetime.now(timezone.utc).isoformat()
        if _history_uses_sheets():
            existing_rows = _sheet_analysis_rows(app_name)
            if any((row.get("run_id") or "").strip() == run_id for row in existing_rows):
                _clear_pending_analysis_state()
                return
            iteration = max(
                [row.get("iteration", 0) for row in existing_rows if (row.get("source_identity") or "") == source_identity] or [0]
            ) + 1
            stored = _sheet_append_row(
                app_name,
                "analysis_runs",
                {
                    "run_id": run_id,
                    "ts_utc": ts_utc,
                    "app": app_name,
                    "email": (email or "").strip().lower(),
                    "source_type": source_type,
                    "source_identity": source_identity,
                    "source_label": source_label,
                    "analysis_key": analysis_key,
                    "iteration": iteration,
                    "spelling_count": int(spelling_count),
                    "grammar_count": int(grammar_count),
                    "editorial_count": int(editorial_count),
                    "fact_count": int(fact_count),
                    "total_count": int(spelling_count + grammar_count + editorial_count + fact_count),
                },
            )
            if not stored:
                raise RuntimeError("sheets-write-failed")
        else:
            ensure_history_db()
            with _history_conn() as conn:
                exists = conn.execute(
                    "SELECT 1 FROM analysis_runs WHERE run_id = ?",
                    (run_id,),
                ).fetchone()
                if exists:
                    _clear_pending_analysis_state()
                    return

                iteration = conn.execute(
                    "SELECT COALESCE(MAX(iteration), 0) + 1 FROM analysis_runs WHERE app = ? AND source_identity = ?",
                    (app_name, source_identity),
                ).fetchone()[0]

                conn.execute(
                    """
                    INSERT INTO analysis_runs (
                        run_id, ts_utc, app, email, source_type, source_identity, source_label,
                        analysis_key, iteration, spelling_count, grammar_count, editorial_count,
                        fact_count, total_count
                    ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                    """,
                    (
                        run_id,
                        ts_utc,
                        app_name,
                        (email or "").strip().lower(),
                        source_type,
                        source_identity,
                        source_label,
                        analysis_key,
                        iteration,
                        int(spelling_count),
                        int(grammar_count),
                        int(editorial_count),
                        int(fact_count),
                        int(spelling_count + grammar_count + editorial_count + fact_count),
                    ),
                )
        _clear_pending_analysis_state()
    except Exception:
        pass

def _fetch_rows(query: str, params=()):
    return _sqlite_rows(query, params)

def render_admin_dashboard(app_name: str):
    st.divider()
    with st.expander("Admin History", expanded=False):
        if _history_uses_sheets():
            login_rows = _sheet_login_rows(app_name)[:200]
            all_analysis_rows = _sheet_analysis_rows(app_name)
            recent_rows = [
                {
                    "ts_utc": row["ts_utc"],
                    "email": row["email"],
                    "source_type": row["source_type"],
                    "source_label": row["source_label"],
                    "iteration": row["iteration"],
                    "spelling_count": row["spelling_count"],
                    "grammar_count": row["grammar_count"],
                    "editorial_count": row["editorial_count"],
                    "fact_count": row["fact_count"],
                    "total_count": row["total_count"],
                }
                for row in all_analysis_rows[:200]
            ]

            daily_map = {}
            for row in login_rows:
                date_utc = (row.get("ts_utc") or "")[:10]
                email = (row.get("email") or "").strip().lower()
                if not date_utc or not email:
                    continue
                entry = daily_map.setdefault(
                    (date_utc, email),
                    {
                        "date_utc": date_utc,
                        "email": email,
                        "login_count": 0,
                        "analyses_run": 0,
                        "distinct_articles": set(),
                        "spelling_issues": 0,
                        "grammar_issues": 0,
                        "editorial_issues": 0,
                        "fact_issues": 0,
                    },
                )
                entry["login_count"] += 1

            for row in all_analysis_rows:
                date_utc = (row.get("ts_utc") or "")[:10]
                email = (row.get("email") or "").strip().lower()
                if not date_utc or not email:
                    continue
                entry = daily_map.setdefault(
                    (date_utc, email),
                    {
                        "date_utc": date_utc,
                        "email": email,
                        "login_count": 0,
                        "analyses_run": 0,
                        "distinct_articles": set(),
                        "spelling_issues": 0,
                        "grammar_issues": 0,
                        "editorial_issues": 0,
                        "fact_issues": 0,
                    },
                )
                entry["analyses_run"] += 1
                entry["distinct_articles"].add((row.get("source_identity") or "").strip())
                entry["spelling_issues"] += _safe_int(row.get("spelling_count"))
                entry["grammar_issues"] += _safe_int(row.get("grammar_count"))
                entry["editorial_issues"] += _safe_int(row.get("editorial_count"))
                entry["fact_issues"] += _safe_int(row.get("fact_count"))

            daily_rows = []
            for entry in daily_map.values():
                daily_rows.append(
                    {
                        "date_utc": entry["date_utc"],
                        "email": entry["email"],
                        "login_count": entry["login_count"],
                        "analyses_run": entry["analyses_run"],
                        "distinct_articles": len([value for value in entry["distinct_articles"] if value]),
                        "spelling_issues": entry["spelling_issues"],
                        "grammar_issues": entry["grammar_issues"],
                        "editorial_issues": entry["editorial_issues"],
                        "fact_issues": entry["fact_issues"],
                    }
                )
            daily_rows.sort(key=lambda row: (row["date_utc"], row["email"]), reverse=True)
            daily_rows = daily_rows[:180]
        else:
            daily_rows = _fetch_rows(
                """
                WITH login_daily AS (
                    SELECT
                        substr(ts_utc, 1, 10) AS date_utc,
                        email,
                        COUNT(*) AS login_count
                    FROM login_events
                    WHERE app = ?
                    GROUP BY substr(ts_utc, 1, 10), email
                ),
                analysis_daily AS (
                    SELECT
                        substr(ts_utc, 1, 10) AS date_utc,
                        email,
                        COUNT(*) AS analyses_run,
                        COUNT(DISTINCT source_identity) AS distinct_articles,
                        COALESCE(SUM(spelling_count), 0) AS spelling_issues,
                        COALESCE(SUM(grammar_count), 0) AS grammar_issues,
                        COALESCE(SUM(editorial_count), 0) AS editorial_issues,
                        COALESCE(SUM(fact_count), 0) AS fact_issues
                    FROM analysis_runs
                    WHERE app = ?
                    GROUP BY substr(ts_utc, 1, 10), email
                ),
                combined AS (
                    SELECT date_utc, email FROM login_daily
                    UNION
                    SELECT date_utc, email FROM analysis_daily
                )
                SELECT
                    c.date_utc,
                    c.email,
                    COALESCE(l.login_count, 0) AS login_count,
                    COALESCE(a.analyses_run, 0) AS analyses_run,
                    COALESCE(a.distinct_articles, 0) AS distinct_articles,
                    COALESCE(a.spelling_issues, 0) AS spelling_issues,
                    COALESCE(a.grammar_issues, 0) AS grammar_issues,
                    COALESCE(a.editorial_issues, 0) AS editorial_issues,
                    COALESCE(a.fact_issues, 0) AS fact_issues
                FROM combined c
                LEFT JOIN login_daily l
                  ON l.date_utc = c.date_utc AND l.email = c.email
                LEFT JOIN analysis_daily a
                  ON a.date_utc = c.date_utc AND a.email = c.email
                ORDER BY c.date_utc DESC, c.email ASC
                LIMIT 180
                """,
                (app_name, app_name),
            )

            login_rows = _fetch_rows(
                """
                SELECT ts_utc, email
                FROM login_events
                WHERE app = ?
                ORDER BY ts_utc DESC
                LIMIT 200
                """,
                (app_name,),
            )

            recent_rows = _fetch_rows(
                """
                SELECT
                    ts_utc,
                    email,
                    source_type,
                    source_label,
                    iteration,
                    spelling_count,
                    grammar_count,
                    editorial_count,
                    fact_count,
                    total_count
                FROM analysis_runs
                WHERE app = ?
                ORDER BY ts_utc DESC
                LIMIT 200
                """,
                (app_name,),
            )

        source_search = st.text_input("Search article or document", key=f"{app_name}_history_search")
        search_rows = []
        if source_search:
            if _history_uses_sheets():
                needle = source_search.strip().lower()
                search_rows = [
                    {
                        "ts_utc": row["ts_utc"],
                        "email": row["email"],
                        "source_type": row["source_type"],
                        "source_label": row["source_label"],
                        "source_identity": row["source_identity"],
                        "iteration": row["iteration"],
                        "spelling_count": row["spelling_count"],
                        "grammar_count": row["grammar_count"],
                        "editorial_count": row["editorial_count"],
                        "fact_count": row["fact_count"],
                        "total_count": row["total_count"],
                    }
                    for row in all_analysis_rows
                    if needle in (row.get("source_label") or "").lower()
                    or needle in (row.get("source_identity") or "").lower()
                ][:200]
            else:
                like = f"%{source_search.strip()}%"
                search_rows = _fetch_rows(
                    """
                    SELECT
                        ts_utc,
                        email,
                        source_type,
                        source_label,
                        source_identity,
                        iteration,
                        spelling_count,
                        grammar_count,
                        editorial_count,
                        fact_count,
                        total_count
                    FROM analysis_runs
                    WHERE app = ?
                      AND (source_label LIKE ? OR source_identity LIKE ?)
                    ORDER BY ts_utc DESC
                    LIMIT 200
                    """,
                    (app_name, like, like),
                )

        st.markdown("#### Daily Summary")
        if daily_rows:
            st.dataframe(daily_rows, use_container_width=True)
        else:
            st.info("No daily history available yet.")

        st.markdown("#### Recent Analysis Runs")
        if recent_rows:
            st.dataframe(recent_rows, use_container_width=True)
        else:
            st.info("No analysis runs recorded yet.")

        st.markdown("#### Recent Logins")
        if login_rows:
            st.dataframe(login_rows, use_container_width=True)
        else:
            st.info("No login history recorded yet.")

        st.markdown("#### Article / Document Iterations")
        if source_search:
            if search_rows:
                st.dataframe(search_rows, use_container_width=True)
            else:
                st.info("No matching article or document history found.")
        else:
            st.caption("Search by URL, filename, or document hash to see iteration history.")

def enforce_app_access(app_title: str, app_caption: str, app_name: str):
    _restore_persisted_session(app_name)
    if _email_access_granted():
        with st.sidebar:
            st.caption(f"Signed in as {st.session_state.get('_email_access_email', '')}")
            if st.button("Log out"):
                _revoke_persisted_session(app_name)
                _clear_email_access()
                st.rerun()
        return

    st.title(app_title)
    st.caption(app_caption)
    with st.form("email_access_login"):
        username = st.text_input("Work email username", placeholder="firstname.lastname")
        st.caption(f"Domain fixed as @{ALLOWED_EMAIL_DOMAIN}")
        submitted = st.form_submit_button("Continue", type="primary")

    st.caption(APP_ACCESS_SUPPORT_TEXT)

    if submitted:
        email = _build_email_from_username(username)
        if not _email_allowed(email):
            st.error("Please enter only your username, without '@' or the domain.")
        else:
            st.session_state["_email_access_granted"] = True
            st.session_state["_email_access_email"] = email
            _record_access_event(app_name, email)
            _create_persisted_session(app_name, email)
            st.rerun()
    st.stop()

if not IMPORT_ONLY:
    enforce_app_access(
        "🧪 Article QC Tool (Gemini 2.5 – Vertex AI)",
        "Spelling · Grammar · Editorial Safety · AI Review",
        "english_qc",
    )
    st.title("🧪 Article QC Tool (Gemini 2.5 – Vertex AI)")
    st.caption("Spelling · Grammar · Editorial Safety · AI Review")


# =================================================
# 🔑 VERTEX AI AUTH (BASE64 SAFE)
# =================================================
PROJECT_ID = str(_secret("VERTEX_PROJECT_ID", "")).strip()
REGION = "us-central1"
CRED_PATH = "/tmp/gcp_service_account.json"
MODEL_PRO = "gemini-2.5-pro"
MODEL_FLASH = "gemini-2.5-flash"
CLOUD_PLATFORM_SCOPE = "https://www.googleapis.com/auth/cloud-platform"
SHEETS_SCOPE = "https://www.googleapis.com/auth/spreadsheets"


def load_gcp_credentials():
    try:
        creds, project_id, _ = _get_scoped_service_account_credentials([CLOUD_PLATFORM_SCOPE])
        return creds, project_id
    except Exception as e:
        st.error("❌ Invalid Base64 GCP credential")
        st.exception(e)
        st.stop()

@st.cache_resource
def init_vertex_and_model():
    creds, project_id = load_gcp_credentials()

    client = genai.Client(
        vertexai=True,
        project=project_id,
        location=REGION,
        credentials=creds,
    )

    # Try pro, fallback to flash
    try:
        client.models.generate_content(
            model=MODEL_PRO,
            contents="Warmup",
            config=genai_types.GenerateContentConfig(
                temperature=0,
                topP=1,
                maxOutputTokens=8,
            ),
        )
        st.success("✅ Gemini 2.5 Pro loaded")
        default_model = MODEL_PRO
    except Exception:
        client.models.generate_content(
            model=MODEL_FLASH,
            contents="Warmup",
            config=genai_types.GenerateContentConfig(
                temperature=0,
                topP=1,
                maxOutputTokens=8,
            ),
        )
        st.warning("⚠️ Falling back to Gemini 2.5 Flash")
        default_model = MODEL_FLASH

    return client, default_model

def build_generate_config(generation_config=None):
    cfg = generation_config or {}
    return genai_types.GenerateContentConfig(
        temperature=cfg.get("temperature"),
        topP=cfg.get("top_p"),
        topK=cfg.get("top_k"),
        candidateCount=cfg.get("candidate_count"),
        maxOutputTokens=cfg.get("max_output_tokens"),
    )

def generate_text(prompt, generation_config=None, model_name=None):
    client, default_model = init_vertex_and_model()
    response = client.models.generate_content(
        model=model_name or default_model,
        contents=prompt,
        config=build_generate_config(generation_config),
    )
    return response.text or ""

def format_ai_error(prefix: str, exc: Exception) -> str:
    return f"__ERROR__:{prefix}: {type(exc).__name__}: {exc}"

def is_ai_error_output(text: str) -> bool:
    value = (text or "").strip()
    return value.startswith("__ERROR__:") or value.startswith("Error:")

def summarise_ai_error(text: str) -> str:
    value = (text or "").replace("__ERROR__:", "").strip()
    lower = value.lower()

    project_match = re.search(r"project[s/ ]+([a-z0-9\\-]+)", value, flags=re.IGNORECASE)
    project_id = project_match.group(1) if project_match else "the configured project"

    if "service_disabled" in lower or "vertex ai api has not been used in project" in lower:
        return (
            f"Vertex AI API is disabled in {project_id}. Enable `aiplatform.googleapis.com`, "
            "wait a few minutes, and retry."
        )
    if "iam_permission_denied" in lower or "aiplatform.endpoints.predict" in lower:
        return (
            f"The configured service account does not have Vertex AI prediction permission "
            f"(`aiplatform.endpoints.predict`) in {project_id}."
        )
    if "invalid jwt signature" in lower or "invalid_grant" in lower:
        return "The configured service-account key is invalid, revoked, or does not match the active secret."
    return value

def render_ai_error(section_label: str, value: str, container=None) -> bool:
    if not is_ai_error_output(value):
        return False
    target = container if container is not None else st
    target.error(f"{section_label}: {summarise_ai_error(value)}")
    return True

def generate_stream_text(prompt, generation_config=None, model_name=None):
    client, default_model = init_vertex_and_model()
    stream = client.models.generate_content_stream(
        model=model_name or default_model,
        contents=prompt,
        config=build_generate_config(generation_config),
    )
    chunks = []
    for chunk in stream:
        if getattr(chunk, "text", None):
            chunks.append(chunk.text)
    return "".join(chunks)


# =================================================
# INPUT EXTRACTION (INLINE-SAFE)
# =================================================
def clean_docx(file_path):
    doc = Document(file_path)
    content, seen = [], set()

    for para in doc.paragraphs:
        txt = para.text.strip()
        if not txt or txt in seen or len(txt) < 15:
            continue
        content.append(("paragraph", txt))
        seen.add(txt)

    return content

def clean_article(url):
    header_profiles = [
        {
            "User-Agent": (
                "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) "
                "AppleWebKit/537.36 (KHTML, like Gecko) "
                "Chrome/137.0.0.0 Safari/537.36"
            ),
            "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
            "Accept-Language": "en-IN,en-GB;q=0.9,en;q=0.8",
            "Cache-Control": "no-cache",
            "Pragma": "no-cache",
        },
        {
            "User-Agent": "Mozilla/5.0",
            "Accept": "*/*",
        },
    ]

    response = None
    last_exc = None

    for headers in header_profiles:
        try:
            response = requests.get(url, headers=headers, timeout=20, allow_redirects=True)
            response.raise_for_status()
            break
        except requests.RequestException as exc:
            last_exc = exc
            response = None
            continue

    if response is None:
        status = ""
        if isinstance(last_exc, requests.HTTPError) and last_exc.response is not None:
            status = f" (HTTP {last_exc.response.status_code})"
        raise RuntimeError(
            "Unable to fetch the article URL from the server"
            f"{status}. Try the DOCX upload option or retry later."
        ) from last_exc

    soup = BeautifulSoup(response.text, "html.parser")

    content = []
    seen = set()

    # ---------- TITLE ----------
    title = soup.find("h1")
    if title:
        content.append(("heading", title.get_text(strip=True)))

    # ---------- TRUE FOOTER STOP MARKERS ----------
    HARD_STOP_MARKERS = [
        "our aim is to provide",
        "explore the world",
        "copyright",
        "all rights reserved",
        "for any feedback",
        "compliant_gro@",
    ]

    # ---------- SOFT SKIP MARKERS (DO NOT BREAK) ----------
    SOFT_SKIP_MARKERS = [
        "don't miss",
        "dont miss",
        "also read",
        "click here",
        "follow us",
    ]

    for el in soup.find_all(["p", "li"], recursive=True):
        txt = el.get_text(separator=" ", strip=True)
        # 🔥 FIX: remove space BEFORE punctuation introduced by HTML
        txt = re.sub(r"\s+([,.;:!?])", r"\1", txt)

        if not txt or len(txt) < 20:
            continue

        lower = txt.lower()

        # 🛑 HARD STOP → footer / legal / site boilerplate
        if any(marker in lower for marker in HARD_STOP_MARKERS):
            break

        # ⛔ SKIP widget headers / navigation
        if any(marker in lower for marker in SOFT_SKIP_MARKERS):
            continue

        # ⛔ SKIP numbered widget lists (1…, 2…, 3…)
        if re.match(r"^\d+\s+", txt):
            continue

        if txt in seen:
            continue

        seen.add(txt)
        content.append(("paragraph", txt))

    return content


# =================================================
# LOAD MODELS
# =================================================
@st.cache_resource
def load_nlp():
    return spacy.load("en_core_web_sm")


@st.cache_resource
def load_languagetool():
    try:
        return language_tool_python.LanguageToolPublicAPI("en-GB")
    except Exception:
        return None


@st.cache_resource
def load_symspell():
    sym = SymSpell(max_dictionary_edit_distance=2, prefix_length=7)
    dictionary_path = str(
        pkg_files("symspellpy").joinpath("frequency_dictionary_en_82_765.txt")
    )
    sym.load_dictionary(dictionary_path, 0, 1)
    return sym


nlp = load_nlp()
lt_tool = load_languagetool()
sym_spell = load_symspell()
spell = SpellChecker()


# =================================================
# SPELLING + GRAMMAR
# =================================================
STOPWORDS = {
    "a","an","the","of","to","in","on","at","for","from","by",
    "and","or","but","so","are","is","was","were","be","been","being"
}

GRAMMAR_MAX_CHARS = 20000
FACT_MAX_CHARS = 20000
FACT_MAX_ITEMS = 40

# =================================================
# EDITORIAL QC RULES (NEW)
# =================================================
US_UK_SPELLINGS = {
    "organize": "organise",
    "organizes": "organises",
    "organized": "organised",
    "organizing": "organising",
    "organization": "organisation",
    "analyze": "analyse",
    "analyzes": "analyses",
    "analyzed": "analysed",
    "analyzing": "analysing",
    "center": "centre",
    "color": "colour",
    "colors": "colours",
    "favor": "favour",
    "favored": "favoured",
    "favorite": "favourite",
    "behavior": "behaviour",
    "defense": "defence",
    "license": "licence",
    "offense": "offence",
    "traveling": "travelling",
    "traveled": "travelled",
    "canceled": "cancelled",
    "counseling": "counselling",
    "program": "programme",
    "gray": "grey",
}

HYPERBOLE_WORDS = [
    "best",
    "awesome",
    "amazing",
    "incredible",
    "mind-blowing",
    "unbelievable",
    "stunning",
    "ultimate",
    "perfect",
    "must-see",
    "must watch",
]

CLICHE_PHRASES = [
    "on the other hand",
    "moreover",
    "in this regard",
    "do the needful",
    "at the end of the day",
    "in a nutshell",
]

URL_RE = re.compile(r"\b(?:https?://|www\.)\S+\b", re.IGNORECASE)
ACRONYM_DOT_RE = re.compile(r"\b(?:[A-Z]\.\s?){2,}\b")
HONORIFIC_RE = re.compile(r"\b(Mr|Mrs|Ms)\.?\b")
HONOURS_PREFIX_RE = re.compile(
    r"\b(Bharat Ratna|Padma (?:Shri|Bhushan|Vibhushan))\s+([A-Z][\w'-]+)\b"
)
AGE_NO_HYPHEN_RE = re.compile(
    r"\b(\d{1,3})\s*(?:year|yr|yrs)\s*old\b", re.IGNORECASE
)
AGE_BAD_HYPHEN_RE = re.compile(r"\b(\d{1,3})-year\s+old\b", re.IGNORECASE)
RANGE_RE = re.compile(r"\b[0-9]\s*[–-]\s*[0-9]\b")
UNIT_RE = re.compile(
    r"\b(\d+(?:\.\d+)?)\s*(kg|kgs|KG|KGS|km|kms|KM|KMS|kmph|KMPH|KMPh|km/h|KM/H)\b"
)
LIST_ENUM_RE = re.compile(r"^\s*\(?\d+[.)]")
LEGAL_CONTEXT_RE = re.compile(
    r"\b(section|sec\.?|article|art\.?|rule|order|clause|schedule|"
    r"sub-?section|sub-?clause|act|ipc|crpc|bnss|bns|cpc)\b",
    re.IGNORECASE
)
LEGAL_REF_RE = re.compile(r"\b\d+\s*\(\d+\)\b")
DATE_TIME_RE = re.compile(
    r"\b(\d{1,2}:\d{2}\s*(?:am|pm)?)\b|\b\d{1,2}\s+[A-Za-z]{3,9}\b",
    re.IGNORECASE
)
DGMENT_RE = re.compile(r"\b[A-Za-z]*dgment(s)?\b")
QUOTE_RE = re.compile(r"[\"“”]")
EXPERT_HINT_RE = re.compile(
    r"\b(Dr|Doctor|Professor|expert|dermatologist|nutritionist|dietitian|doctor|"
    r"physician|surgeon|psychologist|cardiologist|gynecologist|gynaecologist|"
    r"paediatrician|pediatrician|dentist|trichologist)\b",
    re.IGNORECASE
)

HYPERBOLE_RE = re.compile(
    r"\b(" + "|".join(re.escape(w) for w in HYPERBOLE_WORDS) + r")\b",
    re.IGNORECASE
)
CLICHE_RE = re.compile(
    r"\b(" + "|".join(re.escape(p) for p in CLICHE_PHRASES) + r")\b",
    re.IGNORECASE
)

UNIT_NORMALIZATION = {
    "kg": "kg",
    "kgs": "kg",
    "km": "km",
    "kms": "km",
    "kmph": "kmph",
    "km/h": "km/h",
}


def _preserve_case(source, replacement):
    if source.isupper():
        return replacement.upper()
    if source[:1].isupper():
        return replacement.capitalize()
    return replacement


def _excerpt(text, start, end, width=40):
    left = max(0, start - width)
    right = min(len(text), end + width)
    snippet = text[left:right].replace("\n", " ").strip()
    return snippet


def _escape_md(text):
    return text.replace("|", "\\|").replace("\n", " ").strip()


def _render_issues_table(issues):
    if not issues:
        return ""
    lines = [
        "| Rule | Location | Excerpt | Suggestion |",
        "| --- | --- | --- | --- |",
    ]
    for issue in issues:
        rule = _escape_md(issue.get("rule", ""))
        location = _escape_md(issue.get("location", ""))
        excerpt = _escape_md(issue.get("excerpt", ""))[:160]
        suggestion = _escape_md(issue.get("suggestion", ""))
        lines.append(f"| {rule} | {location} | {excerpt} | {suggestion} |")
    return "\n".join(lines)


def _batch_texts(texts, max_chars):
    batches = []
    current = []
    size = 0

    for text in texts:
        t = text.strip()
        if not t:
            continue
        add_len = len(t) + 2
        if current and size + add_len > max_chars:
            batches.append("\n\n".join(current))
            current = [t]
            size = len(t)
        else:
            current.append(t)
            size += add_len

    if current:
        batches.append("\n\n".join(current))

    return batches


def _batch_statements(statements, max_chars, max_items):
    batches = []
    current = []
    size = 0

    for stmt in statements:
        line = f"- {stmt}"
        add_len = len(line) + 1
        if current and (size + add_len > max_chars or len(current) >= max_items):
            batches.append(current)
            current = [stmt]
            size = add_len
        else:
            current.append(stmt)
            size += add_len

    if current:
        batches.append(current)

    return batches


def editorial_qc_checks(content, web_story=False, health_beauty=False):
    issues = []
    seen_issues = set()
    paragraph_index = 0
    found_expert_quote = False

    def add_issue(rule, location, excerpt, suggestion):
        key = (
            rule,
            location,
            re.sub(r"\s+", " ", excerpt.strip().lower())
        )
        if key in seen_issues:
            return
        seen_issues.add(key)
        issues.append({
            "rule": rule,
            "location": location,
            "excerpt": excerpt,
            "suggestion": suggestion,
        })

    for ctype, text in content:
        if not text:
            continue

        if ctype == "paragraph":
            paragraph_index += 1
            location = f"Paragraph {paragraph_index}"
        else:
            location = "Heading"

        doc = nlp(text)
        entity_spans = [(ent.start_char, ent.end_char) for ent in doc.ents]

        def in_entity(start, end):
            return any(start >= s and end <= e for s, e in entity_spans)

        # British English spellings (skip official names/entities)
        for us, uk in US_UK_SPELLINGS.items():
            for match in re.finditer(rf"\b{re.escape(us)}\b", text, re.IGNORECASE):
                if in_entity(match.start(), match.end()):
                    continue
                suggestion = _preserve_case(match.group(), uk)
                add_issue(
                    "British English spelling",
                    location,
                    _excerpt(text, match.start(), match.end()),
                    f"Use '{suggestion}'"
                )

        # British English: -dgment -> -dgement (dynamic, not hardcoded)
        for match in DGMENT_RE.finditer(text):
            if in_entity(match.start(), match.end()):
                continue
            word = match.group(0)
            if "dgement" in word.lower():
                continue
            suggestion = re.sub(r"dgment(s)?$", r"dgement\1", word, flags=re.IGNORECASE)
            suggestion = _preserve_case(word, suggestion)
            add_issue(
                "British English spelling",
                location,
                _excerpt(text, match.start(), match.end()),
                f"Use '{suggestion}'"
            )

        # No ampersands unless part of an official name/entity
        for match in re.finditer(r"&", text):
            if in_entity(match.start(), match.end()):
                continue
            add_issue(
                "No ampersands",
                location,
                _excerpt(text, match.start(), match.end()),
                "Replace '&' with 'and'"
            )

        # Numerals 0-9 in body text (headlines and ranges allowed)
        if ctype == "paragraph":
            range_spans = [(m.start(), m.end()) for m in RANGE_RE.finditer(text)]
            for match in re.finditer(r"\b[0-9]\b", text):
                # Skip list numbering at the start of a line
                if match.start() <= 3 and LIST_ENUM_RE.match(text):
                    continue
                # Skip legal references and section/article numbering
                left = max(0, match.start() - 20)
                right = min(len(text), match.end() + 20)
                context = text[left:right]
                if LEGAL_CONTEXT_RE.search(context) or LEGAL_REF_RE.search(context):
                    continue
                # Skip dates and times
                if DATE_TIME_RE.search(context):
                    continue
                if any(match.start() >= s and match.end() <= e for s, e in range_spans):
                    continue
                add_issue(
                    "Number style (0–9)",
                    location,
                    _excerpt(text, match.start(), match.end()),
                    "Spell out numbers 0–9 in body text"
                )

        # Measurement units: lowercase, singular
        for match in UNIT_RE.finditer(text):
            number = match.group(1)
            unit_raw = match.group(2)
            unit_key = unit_raw.lower()
            if unit_key not in UNIT_NORMALIZATION:
                continue
            corrected_unit = UNIT_NORMALIZATION[unit_key]
            if unit_raw != corrected_unit:
                add_issue(
                    "Unit formatting",
                    location,
                    _excerpt(text, match.start(), match.end()),
                    f"Use '{number} {corrected_unit}'"
                )

        # Acronyms should not include periods/spaces
        for match in ACRONYM_DOT_RE.finditer(text):
            cleaned = re.sub(r"[.\s]", "", match.group())
            add_issue(
                "Acronym punctuation",
                location,
                _excerpt(text, match.start(), match.end()),
                f"Use '{cleaned}'"
            )

        # Hyperlinks: none in first two paragraphs; none at all for web stories
        if URL_RE.search(text):
            if web_story or (ctype == "paragraph" and paragraph_index <= 2):
                add_issue(
                    "Hyperlink placement",
                    location,
                    _excerpt(text, 0, min(len(text), 80)),
                    "Remove hyperlinks (web story)" if web_story
                    else "Avoid external hyperlinks in the first two paragraphs"
                )

        # Hyperbole and cliches
        for match in HYPERBOLE_RE.finditer(text):
            add_issue(
                "Avoid superlatives",
                location,
                _excerpt(text, match.start(), match.end()),
                "Remove or substantiate the claim"
            )
        for match in CLICHE_RE.finditer(text):
            add_issue(
                "Avoid cliches",
                location,
                _excerpt(text, match.start(), match.end()),
                "Rephrase in a fresh, direct way"
            )

        # Honorifics
        for match in HONORIFIC_RE.finditer(text):
            add_issue(
                "Honorifics",
                location,
                _excerpt(text, match.start(), match.end()),
                "Remove Mr/Mrs/Ms"
            )

        # Honours as suffixes
        for match in HONOURS_PREFIX_RE.finditer(text):
            award = match.group(1)
            name = match.group(2)
            add_issue(
                "Honours as suffix",
                location,
                _excerpt(text, match.start(), match.end()),
                f"Use '{award} awardee {name}'"
            )

        # Age formatting
        for match in AGE_BAD_HYPHEN_RE.finditer(text):
            add_issue(
                "Age formatting",
                location,
                _excerpt(text, match.start(), match.end()),
                "Use '38-year-old' or 'Name, 38,'"
            )
        for match in AGE_NO_HYPHEN_RE.finditer(text):
            if re.search(rf"\b{match.group(1)}-year-old\b", text, re.IGNORECASE):
                continue
            add_issue(
                "Age formatting",
                location,
                _excerpt(text, match.start(), match.end()),
                "Use '38-year-old' or 'Name, 38,'"
            )

        # Expert quote detection (health/beauty)
        if health_beauty and QUOTE_RE.search(text) and EXPERT_HINT_RE.search(text):
            found_expert_quote = True

    if health_beauty and not found_expert_quote:
        add_issue(
            "Expert quote required",
            "Overall",
            "No expert quote detected",
            "Add a quote from a relevant expert"
        )

    return issues


def correct_spelling_minimal(text):
    tokens = text.split()
    out = []

    entities = {ent.text for ent in nlp(text).ents}
    entity_tokens = {w for e in entities for w in e.split()}

    for tok in tokens:
        core = re.sub(r"[^\w]", "", tok)
        lower = core.lower()

        if not core or core in entity_tokens or lower in STOPWORDS:
            out.append(tok)
            continue

        suggestions = sym_spell.lookup(lower, Verbosity.CLOSEST, 2)
        best = suggestions[0].term if suggestions else spell.correction(lower)

        if best and best != lower and zipf_frequency(best, "en") > zipf_frequency(lower, "en"):
            best = best.capitalize() if core[0].isupper() else best
            out.append(tok.replace(core, best))
        else:
            out.append(tok)

    return " ".join(out)


def correct_grammar_languagetool(text):
    if not lt_tool:
        return text

    try:
        return language_tool_python.utils.correct(text, lt_tool.check(text))
    except Exception:
        return text


# =================================================
# 🔍 VERBATIM DIFF VALIDATOR + CONFIDENCE SCORE
# =================================================
def filter_gemini_rows(raw_table, article_text):
    lines = raw_table.splitlines()
    output = []

    header_added = False

    for line in lines:
        if line.strip().startswith("| Original"):
            output.append("| Original | Corrected | Reason |")
            output.append("|---|---|---|")
            header_added = True
            continue

        if "|" not in line:
            continue

        cols = [c.strip() for c in line.split("|") if c.strip()]
        if len(cols) != 3:
            continue

        original, corrected, reason = cols

        # Verbatim safety: original must exist exactly in article text
        if original not in article_text:
            continue

        if is_noop_correction(original, corrected, reason):
            continue

        output.append(f"| {original} | {corrected} | {reason} |")

    return "\n".join(output) if header_added else ""

# =================================================
# GEMINI QC — OLD TABLE FORMAT (SAFE)
# =================================================
def gemini_grammar_review(article_data, max_chars=GRAMMAR_MAX_CHARS):
    init_vertex_and_model()  # ensures vertexai.init() is called

    raw_paragraphs = [
        text
        for ctype, text in article_data
        if ctype == "paragraph" and len(text.split()) >= 6
    ]

    if not raw_paragraphs:
        return ""

    joined = "\n\n".join(raw_paragraphs)
    if len(joined) <= max_chars:
        paragraphs = [joined]
    else:
        paragraphs = _batch_texts(raw_paragraphs, max_chars)

    BASE_PROMPT = """
You are a professional proofreader and a content QC professional.

You are a professional proofreader and a content QC professional.

For maximum editorial efficiency, strictly apply the following stylistic directives:

STYLE & FORMATTING DIRECTIVES (MANDATORY):
- Use the MM DD, YYYY format for all dates.
- Eliminate periods in names, academic degrees, and titles
  (e.g., APJ Abdul Kalam, MSc, Dr — NOT A.P.J. Abdul Kalam, M.Sc., Dr.).
- Spell out the full name at first mention followed by the abbreviation;
  use the abbreviation for all subsequent mentions
  (e.g., State Bank of India, then SBI).
- Write numbers zero through nine alphabetically.
  Use numerals for 10 and above,
  EXCEPT for currency, percentages, time, dates, and ranges.
- Use the full name at first instance.
  Subsequently:
  - Use the surname only when accompanied by a title (Dr, Prof, Lt, Col).
  - If no title exists, use the first name
    (surname only for widely recognised senior figures).
- Avoid transitional words such as "meanwhile" and "moreover".
- Do NOT use em-dashes.
- Always apply the Oxford comma in lists of three or more items.
- Minimise repetitive phrasing or redundancy.
  Avoid repeated possessive pronouns in lists.
- Enclose titles of movies, songs, plays, books, and usernames in single quotes
  (e.g., 'DDLJ').
- Use double quotes exclusively for direct statements and verbatim quotations.
- Use lowercase "am" and "pm" for all time references.
Rules (STRICT):
- Review each paragraph independently
- Only fix spelling, grammar, and language-standard issues
- Do NOT change numbers or numerical values
- British English is the ONLY accepted standard
- Convert American English spellings AND formats to British English where applicable
- British date format must be used (e.g., "1 January", not "January 1")
- Date-order corrections are allowed ONLY when the numeric value remains unchanged
- Understand context before suggesting corrections
- If an issue exists once, it MUST be reported every time
- List *every* spelling correction independently, even if it seems minor

PROHIBITIONS:
- NEVER change proper nouns, political parties, or person names
- NEVER rename quoted speakers
- NEVER modify social media platform names or product/platform identifiers
  (e.g., X, Twitter, Facebook, Instagram)
- NEVER modify single-letter proper nouns (e.g., X)
- Do NOT hallucinate
- Do NOT normalize legal, political, or platform references

CRITICAL CONSTRAINTS:
- You may ONLY use text that appears verbatim in the TEXT section
- NEVER invent new examples, phrases, or sentences
- The "Original" column MUST be an exact, character-for-character substring
  of the provided TEXT
- If no correction is required, DO NOT create a table row
- If you cannot find an exact match in the TEXT, do NOT include it

ABSOLUTE RULE:
- Treat the TEXT as a raw byte string
- Do NOT normalize whitespace beyond the correction itself
- Periods, commas, apostrophes, abbreviations, and numerals must be preserved exactly

ABBREVIATION SAFETY:
- Single-letter abbreviations followed by a period (e.g., "S.", "X.") are VALID

INLINE CONTENT SAFETY:
- Hyperlinks or anchor text may exist
- Treat input as already-rendered plain text
- Do NOT infer missing spaces caused by HTML

PLATFORM NAME SAFETY:
- The platform "X" must NEVER be interpreted as "A"

EXHAUSTIVENESS REQUIREMENT (MANDATORY):
- You MUST scan the entire TEXT from start to end
- You MUST identify ALL applicable issues before responding
- Do NOT prioritise or skip issues due to importance
- Do NOT stop early once some issues are found
- The output must be COMPLETE and repeatable for the same TEXT

Return output strictly as a table:
| Original | Corrected | Reason |
"""

    responses = []

    def call_gemini(prompt):
        return generate_text(
            prompt,
            generation_config={
                "temperature": 0,
                "top_p": 1,
                "top_k": 1,
                "candidate_count": 1
            },
            model_name=MODEL_FLASH,
        )

    # Parallel batch calls for speed (same logic/output)
    with ThreadPoolExecutor(max_workers=3) as executor:
        futures = []
        for para in paragraphs:
            prompt = BASE_PROMPT + "\n\nTEXT:\n" + para
            futures.append(executor.submit(call_gemini, prompt))

        for future in as_completed(futures):
            try:
                responses.append(future.result())
            except Exception:
                continue

    if not responses:
        return ""

    raw = "\n".join(responses)

    # 🔧 Existing extraction + dedupe (UNCHANGED)
    matches = re.findall(
        r"\|\s*([^|]+?)\s*\|\s*([^|]+?)\s*\|\s*([^|]+?)\s*\|",
        raw
    )

    rows = []
    seen = set()

    def canon(s: str) -> str:
        s = s.lower().strip()
        s = re.sub(r"\s+", " ", s)
        return s.strip(".,;:!?")

    for original, corrected, reason in matches:
        if original.strip().lower() == "original":
            continue

        key = (canon(original), canon(corrected), canon(reason))
        if key in seen:
            continue

        seen.add(key)
        rows.append(
            f"| {original.strip()} | {corrected.strip()} | {reason.strip()} |"
        )

    if not rows:
        return ""

    return "\n".join([
        "| Original | Corrected | Reason |",
        "|---|---|---|",
        *rows
    ])


# =================================================
# GEMINI EDITORIAL QC — GUIDELINES
# =================================================
def gemini_editorial_review(article_data, web_story=False, health_beauty=False):
    init_vertex_and_model()

    paragraphs = [
        text[:900]
        for ctype, text in article_data
        if ctype == "paragraph"
    ]

    joined_paragraphs = "\n\n".join(paragraphs)

    context_notes = []
    if web_story:
        context_notes.append("- This is a web story: no hyperlinks are allowed.")
    if health_beauty:
        context_notes.append(
            "- This is a health/beauty story: include at least one expert quote."
        )
    extra_context = "\n".join(context_notes)

    prompt = f"""
You are an editorial QC reviewer. Identify only clear violations and keep fixes concise.

Rules:
1) British English spellings only (e.g., organise). Keep official names as-is.
2) No ampersands unless part of an official title.
3) Single quotes for titles; double quotes only for direct speech or press releases.
4) Use Oxford comma where needed for clarity.
5) Units: lowercase and singular (kg, km, kmph/km/h). No "kgs"/"kms".
6) Acronyms without periods or spaces (US, not U.S.).
7) No external hyperlinks in the first two paragraphs. Web stories: no hyperlinks at all.
8) Avoid superlatives/hyperbole and cliches. Avoid redundancy.
9) Seasons: use singular terms; "rains" only for the entire season.
10) Honorifics: no Mr/Mrs/Ms. Dr/Prof acceptable for doctors/professors.
11) Honours as suffixes (e.g., "Bharat Ratna awardee [Name]").
12) Age format: "[Name], 38," or "38-year-old [Name]".
13) Use man/woman for 18+, boy/girl for 17 and under; minor/juvenile only in legal context.
14) Health/beauty stories must include an expert quote.
15) Provide due credit for data and avoid plagiarism.
16) AI tools may assist but must not be the primary writing source.

Context:
{extra_context if extra_context else "- No extra context."}

Return output strictly as a markdown table:
| Issue | Location | Suggestion |

If there are no issues, return exactly one row:
| No issues found | - | - |

TEXT:
{joined_paragraphs}
"""

    return generate_text(
        prompt,
        generation_config={
            "temperature": 0,
            "top_p": 1,
            "top_k": 1,
            "candidate_count": 1
        },
        model_name=MODEL_FLASH,
    )


# =================================================
# CACHED GEMINI WRAPPERS (STABLE OUTPUTS)
# =================================================
@st.cache_data(show_spinner=False)
def cached_gemini_grammar_review(article_data, max_chars=GRAMMAR_MAX_CHARS):
    return gemini_grammar_review(article_data, max_chars)


@st.cache_data(show_spinner=False)
def cached_gemini_editorial_review(article_data, web_story=False, health_beauty=False):
    return gemini_editorial_review(article_data, web_story, health_beauty)


@st.cache_data(show_spinner=False)
def cached_gemini_fact_check(article_data, max_chars=FACT_MAX_CHARS, max_items=FACT_MAX_ITEMS):
    return gemini_fact_check(article_data, max_chars, max_items)


# ============================
# No-op correction filters
# ============================
def normalize_for_equality(text: str) -> str:
    text = html.unescape((text or "").strip())
    text = text.replace("\u00a0", " ")
    text = re.sub(r"\s+", " ", text)
    return text


def is_noop_reason(reason: str) -> bool:
    normalised = normalize_for_equality(reason).lower()
    return normalised in {
        "",
        "no correction needed",
        "no corrections needed",
        "no change needed",
        "no changes needed",
        "already correct",
        "correct as is",
        "no issue",
        "no issues",
        "no error",
        "no errors",
    }


def is_noop_correction(original: str, corrected: str, reason: str = "") -> bool:
    if is_noop_reason(reason):
        return True
    return normalize_for_equality(original) == normalize_for_equality(corrected)


# ============================
# Invalid rows
# ============================
def filter_invalid_rows(gemini_md, article_text):
    lines = gemini_md.splitlines()
    out = []
    seen = set()

    def normalise(text):
        return re.sub(r"[^a-z0-9]", "", text.lower())

    for line in lines:
        if "|" not in line or line.strip().startswith("| Original"):
            out.append(line)
            continue

        cols = [c.strip() for c in line.split("|") if c.strip()]
        if len(cols) != 3:
            continue

        original, corrected, reason = cols

        # must exist verbatim
        if original not in article_text:
            continue

        if is_noop_correction(original, corrected, reason):
            continue

        # 🔥 FIX #1: normalised edit signature (stable dedupe)
        key = (
            normalise(original),
            normalise(corrected),
        )

        if key in seen:
            continue

        seen.add(key)
        out.append(f"| {original} | {corrected} | {reason} |")

    return "\n".join(out)

# ==============================================
# Split spelling Grammar Function
# ==============================================

def split_spelling_grammar(table_md: str):
    spelling_rows = []
    grammar_rows = []

    rows = re.findall(
        r"\|\s*([^|]+?)\s*\|\s*([^|]+?)\s*\|\s*([^|]+?)\s*\|",
        table_md
    )

    for original, corrected, reason in rows:
        if original.lower() == "original":
            continue

        if is_noop_correction(original, corrected, reason):
            continue

        # 🔑 MINIMAL FIX
        is_spelling = (
            len(original.split()) == 1
            and len(corrected.split()) == 1
            and original != corrected
        )


        row = f"| {original} | {corrected} | {reason} |"

        if is_spelling:
            spelling_rows.append(row)
        else:
            grammar_rows.append(row)

    def sort_key(row):
        cols = [c.strip() for c in row.split("|") if c.strip()]
        if len(cols) != 3:
            return ("", "", "")
        return tuple(re.sub(r"\W+", "", c.lower()) for c in cols)

    spelling_rows.sort(key=sort_key)
    grammar_rows.sort(key=sort_key)

    def build_table(rows):
        if not rows:
            return ""
        return "\n".join(
            ["| Original | Corrected | Reason |",
             "|---|---|---|"] + rows
        )

    return build_table(spelling_rows), build_table(grammar_rows)


def tokenize_for_diff(text: str):
    return re.findall(r"\s+|[\w]+|[^\w\s]", text or "", flags=re.UNICODE)


def highlight_diff_pair(original: str, corrected: str):
    original_tokens = tokenize_for_diff(original)
    corrected_tokens = tokenize_for_diff(corrected)
    matcher = SequenceMatcher(a=original_tokens, b=corrected_tokens)

    original_parts = []
    corrected_parts = []

    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        original_chunk = "".join(original_tokens[i1:i2])
        corrected_chunk = "".join(corrected_tokens[j1:j2])

        if tag == "equal":
            original_parts.append(html.escape(original_chunk))
            corrected_parts.append(html.escape(corrected_chunk))
            continue

        if original_chunk:
            original_parts.append(
                f'<span class="qc-diff qc-diff-original">{html.escape(original_chunk)}</span>'
            )
        if corrected_chunk:
            corrected_parts.append(
                f'<span class="qc-diff qc-diff-corrected">{html.escape(corrected_chunk)}</span>'
            )

    return "".join(original_parts), "".join(corrected_parts)


def parse_language_table_md(table_md: str):
    rows = []
    matches = re.findall(
        r"\|\s*([^|]+?)\s*\|\s*([^|]+?)\s*\|\s*([^|]+?)\s*\|",
        table_md or "",
    )

    for original, corrected, reason in matches:
        if original.lower() == "original":
            continue
        rows.append((original.strip(), corrected.strip(), reason.strip()))

    return rows


def render_language_table(table_md: str):
    rows = parse_language_table_md(table_md)
    if not rows:
        return ""

    lines = [
        """
<style>
.qc-table {
  width: 100%;
  border-collapse: collapse;
  margin-bottom: 1rem;
}
.qc-table th, .qc-table td {
  border: 1px solid rgba(250,250,250,0.14);
  padding: 0.75rem 0.9rem;
  vertical-align: top;
  text-align: left;
}
.qc-table .qc-diff-original {
  color: #ff6b6b;
  font-weight: 700;
}
.qc-table .qc-diff-corrected {
  color: #4ade80;
  font-weight: 700;
}
</style>
<table class="qc-table">
  <thead>
    <tr>
      <th>Original</th>
      <th>Corrected</th>
      <th>Reason</th>
    </tr>
  </thead>
  <tbody>
        """.strip()
    ]

    for original, corrected, reason in rows:
        original_html, corrected_html = highlight_diff_pair(original, corrected)
        lines.append(
            "<tr>"
            f"<td>{original_html}</td>"
            f"<td>{corrected_html}</td>"
            f"<td>{html.escape(reason)}</td>"
            "</tr>"
        )

    lines.append("</tbody></table>")
    return "\n".join(lines)

def parse_markdown_table_rows(table_md: str, expected_columns: int):
    rows = []
    for line in (table_md or "").splitlines():
        row = line.strip()
        if not row.startswith("|") or row.count("|") < expected_columns:
            continue
        parts = [part.strip() for part in row.strip("|").split("|")]
        if len(parts) != expected_columns:
            continue
        if all(re.fullmatch(r":?-{2,}:?", part or "") for part in parts):
            continue
        if parts[0].lower() in {"original", "issue", "statement"}:
            continue
        rows.append(parts)
    return rows

def build_english_qc_report_pdf(source_label: str, user_email: str, spelling_md: str, grammar_md: str, editorial_md: str, fact_md: str):
    try:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    except Exception:
        return None, "PDF export requires the `reportlab` package."

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "qc-title",
        parent=styles["Title"],
        fontName="Helvetica-Bold",
        fontSize=18,
        leading=22,
        textColor=colors.HexColor("#111827"),
    )
    body_style = ParagraphStyle(
        "qc-body",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=9,
        leading=12,
        textColor=colors.HexColor("#111827"),
    )
    heading_style = ParagraphStyle(
        "qc-heading",
        parent=styles["Heading2"],
        fontName="Helvetica-Bold",
        fontSize=12,
        leading=15,
        textColor=colors.HexColor("#1f2937"),
        spaceBefore=10,
        spaceAfter=6,
    )

    def p(text):
        safe = html.escape((text or "").replace("\n", " "))
        return Paragraph(safe, body_style)

    def add_table(story, title, headers, rows, widths):
        story.append(Paragraph(title, heading_style))
        if not rows:
            story.append(Paragraph("No issues found", body_style))
            story.append(Spacer(1, 0.2 * cm))
            return

        data = [[Paragraph(html.escape(h), body_style) for h in headers]]
        for row in rows:
            data.append([p(cell) for cell in row])

        table = Table(data, colWidths=widths, repeatRows=1)
        table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#E5E7EB")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#111827")),
            ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#D1D5DB")),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("LEFTPADDING", (0, 0), (-1, -1), 6),
            ("RIGHTPADDING", (0, 0), (-1, -1), 6),
            ("TOPPADDING", (0, 0), (-1, -1), 5),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
        ]))
        story.append(table)
        story.append(Spacer(1, 0.25 * cm))

    spelling_rows = parse_markdown_table_rows(spelling_md, 3)
    grammar_rows = parse_markdown_table_rows(grammar_md, 3)
    editorial_rows = parse_markdown_table_rows(editorial_md, 3)
    fact_rows = parse_markdown_table_rows(fact_md, 3)

    summary_rows = [
        ["Source", source_label or "-"],
        ["User", user_email or "-"],
        ["Generated (UTC)", datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M:%S")],
        ["Spelling issues", str(len(spelling_rows))],
        ["Grammar issues", str(len(grammar_rows))],
        ["Editorial issues", str(len(editorial_rows))],
        ["Fact check issues", str(len(fact_rows))],
    ]

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=1.2 * cm,
        leftMargin=1.2 * cm,
        topMargin=1.2 * cm,
        bottomMargin=1.2 * cm,
    )

    story = [
        Paragraph("English QC Report", title_style),
        Spacer(1, 0.25 * cm),
    ]

    add_table(story, "Summary", ["Field", "Value"], summary_rows, [4.2 * cm, 12.4 * cm])
    add_table(story, "Spelling Issues", ["Original", "Corrected", "Reason"], spelling_rows, [6.0 * cm, 6.0 * cm, 4.6 * cm])
    add_table(story, "Grammar Issues", ["Original", "Corrected", "Reason"], grammar_rows, [6.0 * cm, 6.0 * cm, 4.6 * cm])
    add_table(story, "Gemini Editorial Review", ["Issue", "Location", "Suggestion"], editorial_rows, [4.6 * cm, 3.0 * cm, 9.0 * cm])
    add_table(story, "Fact Check", ["Statement", "Issue", "Correct Fact"], fact_rows, [6.0 * cm, 4.0 * cm, 6.6 * cm])

    doc.build(story)
    return buffer.getvalue(), None

# =============================
# Is structural Inline
# =============================
def is_structural_line(text: str) -> bool:
    t = text.strip().lower()

    # Headings / labels
    if len(t.split()) <= 3:
        return True

    # Section headers
    if any(t.startswith(x) for x in [
        "day ",
        "note:",
        "about ",
        "directed by",
        "cast",
    ]):
        return True

    # Lists / bullets
    if re.match(r"^\d+[\).]", t):
        return True

    return False

# =================================================
# FACT CHECK — STATEMENT EXTRACTION (DETERMINISTIC)
# =================================================
FACTCHECK_HEADLINE_MARKERS_EN = (
    "fact check",
    "fact-check",
)

FACTCHECK_CONCLUSION_MARKERS_EN = (
    "conclusion:",
    "the claim is false",
    "the claim is misleading",
    "the claim is fake",
    "this claim is false",
    "this claim is misleading",
    "there is no evidence",
    "is unrelated",
    "is not related",
)

FACT_PROCESS_MARKERS_EN = (
    "archive link",
    "archived link",
    "shared on",
    "posted on",
    "posted this",
    "shared this",
    "youtube channel",
    "google lens",
    "keyword search",
    "we searched",
    "we found",
    "we checked",
    "we contacted",
    "fact checker",
    "profile scan",
    "profile was scanned",
    "click here",
    "read here",
    "uploaded on",
    "uploaded to",
    "instagram user",
    "x user",
    "facebook user",
)

def is_english_fact_check_article(article_data):
    for ctype, text in article_data or []:
        if ctype != "heading":
            continue
        lower = (text or "").strip().lower()
        if any(marker in lower for marker in FACTCHECK_HEADLINE_MARKERS_EN):
            return True

    return any(
        any(marker in (text or "").lower() for marker in FACTCHECK_CONCLUSION_MARKERS_EN)
        for ctype, text in (article_data or [])
        if ctype == "paragraph"
    )

def is_english_fact_process_sentence(sentence: str) -> bool:
    lower = (sentence or "").strip().lower()
    return any(marker in lower for marker in FACT_PROCESS_MARKERS_EN)

def is_low_value_fact_sentence_en(sentence: str) -> bool:
    compact = re.sub(r"\s+", " ", (sentence or "").strip())
    lower = compact.lower()
    if re.fullmatch(r"(digital desk,?\s+)?[A-Za-z .'-]+,\s+[A-Za-z .'-]+", compact):
        return True
    if len(lower.split()) < 6:
        return True
    return False

def is_material_fact_candidate_en(sentence: str) -> bool:
    s = (sentence or "").strip()
    if len(s) < 35:
        return False
    if re.search(r"\d", s):
        return True
    return bool(re.search(
        r"\b(is|was|are|were|has|have|had|will|announced|launched|reported|said|claims|according|warned|confirmed|issued)\b",
        s.lower()
    ))

def extract_fact_statements(article_data):
    """
    Deterministically extract candidate factual statements.
    SAME input → SAME statements → EVERY iteration.
    """
    statements = []
    seen = set()
    fact_check_mode = is_english_fact_check_article(article_data)
    lead_paragraphs_taken = 0

    def add_statement(s: str):
        key = re.sub(r"\s+", " ", s.lower())
        if key in seen:
            return
        seen.add(key)
        statements.append(s)

    for ctype, text in article_data:
        if ctype not in {"heading", "paragraph", "table"}:
            continue

        if fact_check_mode:
            if ctype == "heading":
                lower_heading = (text or "").strip().lower()
                if any(marker in lower_heading for marker in FACTCHECK_HEADLINE_MARKERS_EN):
                    add_statement(text.strip())
                continue

            if is_english_fact_process_sentence(text):
                continue

            if any(marker in (text or "").lower() for marker in FACTCHECK_CONCLUSION_MARKERS_EN):
                doc = nlp(text)
                for sent in doc.sents:
                    s = sent.text.strip()
                    if is_english_fact_process_sentence(s) or is_low_value_fact_sentence_en(s):
                        continue
                    if is_material_fact_candidate_en(s):
                        add_statement(s)
                continue

            if ctype == "paragraph" and lead_paragraphs_taken < 3:
                lead_paragraphs_taken += 1
                doc = nlp(text)
                for sent in doc.sents:
                    s = sent.text.strip()
                    if is_english_fact_process_sentence(s) or is_low_value_fact_sentence_en(s):
                        continue
                    if is_material_fact_candidate_en(s):
                        add_statement(s)
            continue

        doc = nlp(text)
        for sent in doc.sents:
            s = sent.text.strip()

            # Basic factual heuristic (NO hard stops, NO assumptions)
            if is_low_value_fact_sentence_en(s):
                continue

            if is_english_fact_process_sentence(s):
                continue

            if not is_material_fact_candidate_en(s):
                continue

            add_statement(s)

    return statements

# =============
# Def Chunked
# =============

def chunked(lst, size):
    for i in range(0, len(lst), size):
        yield lst[i:i + size]

def is_no_issue_fact(issue: str, correction: str) -> bool:
    issue_lower = (issue or "").strip().lower()
    correction_lower = (correction or "").strip().lower()
    return issue_lower in {"", "-", "--", "---", "no issue", "no issues"} or correction_lower in {"", "-", "--", "---", "no issue", "no issues"}

def is_style_only_fact(statement: str, issue: str, correction: str) -> bool:
    lower_issue = (issue or "").strip().lower()
    lower_correction = (correction or "").strip().lower()
    combined = f"{lower_issue} {lower_correction}"
    if re.sub(r"\W+", "", (statement or "").lower()) == re.sub(r"\W+", "", (correction or "").lower()):
        return True
    style_markers = (
        "spelling",
        "grammar",
        "punctuation",
        "style",
        "format",
        "wording",
        "terminology",
        "abbreviation",
        "full form",
        "quote",
        "comma",
        "should be",
        "use ",
        "replace ",
        "preferred",
    )
    return any(marker in combined for marker in style_markers)

def is_current_verification_issue(issue: str) -> bool:
    lower = (issue or "").strip().lower()
    return "needs verification" in lower and "current" in lower

def normalize_fact_correction(issue: str, correction: str, today_iso: str) -> str:
    if is_current_verification_issue(issue):
        return f"Could not verify reliably as of {today_iso}."
    return (correction or "").strip()

# =================================================
# FACT CHECK — SECOND PASS (FAST, STREAMING, STABLE)
# =================================================
def gemini_fact_check(article_data, max_chars=FACT_MAX_CHARS, max_items=FACT_MAX_ITEMS):
    client, _ = init_vertex_and_model()

    # 1️⃣ Deterministic statement universe
    statements = extract_fact_statements(article_data)
    if not statements:
        return ""

    # Full article text (verbatim, unchanged)
    full_text = "\n".join(
        text for ctype, text in article_data if ctype in {"heading", "paragraph", "table"}
    )

    batches = _batch_statements(statements, max_chars, max_items)

    rows = []
    seen = set()
    had_success = False
    last_error = None
    today_iso = datetime.now(timezone.utc).date().isoformat()

    def call_batch(batch):
        batch_block = "\n".join(f"- {stmt}" for stmt in batch)

        fact_prompt = f"""
You are a factual accuracy reviewer for current news copy.

TODAY'S DATE:
{today_iso}

SCOPE:
- Use Google Search grounding for up-to-date verification
- Only evaluate statements that appear verbatim in the TEXT
- Quote the EXACT sentence fragment under "Statement"
- Do NOT paraphrase, rewrite, or infer

EVALUATION RULES:
- For present-tense or current-status claims, verify using information available as of today's date
- For explicitly dated historical claims, judge them against the date or period stated in the article
- If a statement is likely false, mark Issue as "Likely false" and provide the correct fact
- Use "Needs verification (current)" only as a last resort for a materially important current-affairs claim whose present status cannot be verified reliably
- If grounded search is merely sparse, mixed, or not clearly authoritative for a minor claim, omit the row instead of emitting "Needs verification (current)"
- If a statement is likely true, omit it (do NOT create a row)
- NEVER invent facts; if unsure, use "Needs verification (current)"
- Never rely on stale model memory when grounded search is missing or disagrees
- If you do use "Needs verification (current)", the Correct Fact must be exactly: "Could not verify reliably as of {today_iso}."

DO NOT:
- Check grammar, spelling, or style
- Rewrite sentences

Return output strictly as a table:
| Statement | Issue | Correct Fact |

TEXT:
{full_text}

        STATEMENTS:
{batch_block}
"""

        try:
            response = client.models.generate_content(
                model=MODEL_FLASH,
                contents=fact_prompt,
                config=genai_types.GenerateContentConfig(
                    temperature=0,
                    topP=1,
                    topK=1,
                    candidateCount=1,
                    maxOutputTokens=768,
                    seed=0,
                    responseMimeType="text/plain",
                    tools=[genai_types.Tool(google_search=genai_types.GoogleSearch())],
                ),
            )
            return response.text or "", None
        except Exception as exc:
            return "", exc

    batch_results = []
    for batch in batches:
        out, exc = call_batch(batch)
        if exc is not None:
            last_error = exc
            continue
        had_success = True
        batch_results.append(out)

    for out in batch_results:
        if not out:
            continue

        # 3️⃣ Extract table rows (UNCHANGED)
        matches = re.findall(
            r"\|\s*([^|]+?)\s*\|\s*([^|]+?)\s*\|\s*([^|]+?)\s*\|",
            out
        )

        for s, issue, correction in matches:
            if s.lower() == "statement":
                continue
            if any(x.strip() in {"-", "--", "---"} for x in (s, issue, correction)):
                continue
            if is_no_issue_fact(issue, correction):
                continue
            if is_style_only_fact(s, issue, correction):
                continue

            correction = normalize_fact_correction(issue, correction, today_iso)

            sig = (
                re.sub(r"\W+", "", s.lower()),
                re.sub(r"\W+", "", issue.lower())
            )

            if sig in seen:
                continue

            seen.add(sig)
            rows.append((s.strip(), issue.strip(), correction.strip()))

    if not rows and not had_success:
        return format_ai_error("fact", last_error or RuntimeError("No Gemini response"))

    if not rows:
        return ""

    def sort_key(row):
        return (
            re.sub(r"\W+", "", row[0].lower()),
            re.sub(r"\W+", "", row[1].lower()),
            re.sub(r"\W+", "", row[2].lower()),
        )

    rows.sort(key=sort_key)

    rendered = [
        f"| {s} | {issue} | {correction} |"
        for s, issue, correction in rows
    ]

    return "\n".join([
        "| Statement | Issue | Correct Fact |",
        "|---|---|---|",
        *rendered
    ])


# =================================================
# PIPELINE
# =================================================
def run_pipeline(content):
    final = []
    for ctype, text in content:
        if ctype != "paragraph":
            final.append((ctype, text))
            continue

        # ❗ DO NOT mutate article text before Gemini
        final.append((ctype, text))

    return final

# =================================================
# STREAMLIT UI
# =================================================
if not IMPORT_ONLY:
    source = st.sidebar.radio("Source", ["URL", "DOCX"])
    
    st.sidebar.header("QC Options")
    web_story = st.sidebar.checkbox("Web story (no hyperlinks)", value=False)
    health_beauty = st.sidebar.checkbox(
        "Health/Beauty story (expert quote required)", value=False
    )
    run_gemini_editorial = st.sidebar.checkbox(
        "Run Gemini editorial QC", value=True
    )
    if st.sidebar.button("Clear cached AI outputs"):
        st.cache_data.clear()
        _clear_pending_analysis_state()
        for key in (
            "analysis_results",
            "analysis_key",
            "analysis_start",
            "article_content",
            "input_key",
            "source_label",
        ):
            st.session_state.pop(key, None)
    
    analyze_clicked = st.sidebar.button("Analyze")
    
    article_content = None
    current_key = None
    
    if source == "URL":
        url = st.sidebar.text_input("Article URL")
        if url:
            current_key = f"url:{url.strip()}"
        if analyze_clicked and url:
            try:
                article_content = clean_article(url)
                st.session_state["article_content"] = article_content
                st.session_state["input_key"] = current_key
                st.session_state["source_label"] = url.strip()
                queue_analysis_run("url", current_key, url.strip())
            except Exception as exc:
                st.error(str(exc))
                article_content = None
                _clear_pending_analysis_state()
                st.session_state.pop("article_content", None)
                st.session_state.pop("input_key", None)
    else:
        uploaded = st.sidebar.file_uploader("Upload DOCX", type=["docx"])
        if uploaded:
            file_bytes = uploaded.getvalue()
            current_key = "docx:" + hashlib.sha256(file_bytes).hexdigest()
            if analyze_clicked:
                with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as f:
                    f.write(file_bytes)
                    article_content = clean_docx(f.name)
                st.session_state["article_content"] = article_content
                st.session_state["input_key"] = current_key
                st.session_state["source_label"] = uploaded.name or current_key
                queue_analysis_run("docx", current_key, uploaded.name or current_key)
    
    if article_content is None:
        if current_key and st.session_state.get("input_key") == current_key:
            article_content = st.session_state.get("article_content")
    
    source_label = st.session_state.get("source_label", "")
    
    analysis_key = None
    if current_key:
        analysis_key = hashlib.sha256(
            f"{current_key}|{web_story}|{health_beauty}|{run_gemini_editorial}".encode("utf-8")
        ).hexdigest()
    
    analysis_ready = False
    if article_content and analysis_key:
        if analyze_clicked:
            if st.session_state.get("analysis_key") != analysis_key:
                st.session_state["analysis_key"] = analysis_key
                st.session_state["analysis_results"] = {}
                st.session_state["analysis_start"] = time.perf_counter()
            else:
                st.info("Using cached results. Clear cached AI outputs to refresh.")
        if st.session_state.get("analysis_key") == analysis_key:
            analysis_ready = True
        else:
            st.info("Input or options changed. Click Analyze to refresh.")
    
    if analysis_ready:
        qc_content = run_pipeline(article_content)
        results = st.session_state.setdefault("analysis_results", {})
        if "analysis_start" not in results:
            results["analysis_start"] = st.session_state.get("analysis_start", time.perf_counter())
        report_pdf_bytes = None
        report_pdf_error = None
    
        # ---------- FINAL ARTICLE ----------
        st.subheader("📄 Final Article")
        for _, t in qc_content:
            st.write(t)
    
        st.divider()
    
        # ---------- GEMINI QC ----------
        st.subheader("🤖 Gemini QC Review")
        score_placeholder = st.empty()
    
        article_text = "\n".join(
            t for c, t in article_content if c == "paragraph"
        )
        # Grammar + Spelling placeholders
        st.markdown("### ✍️ Spelling Issues")
        spelling_placeholder = st.empty()
        st.markdown("### 🧠 Grammar Issues")
        grammar_placeholder = st.empty()
    
        editorial_placeholder = None
        if run_gemini_editorial:
            st.markdown("### 🧠 Gemini Editorial Review")
            editorial_placeholder = st.empty()
    
        # ---------- FACT CHECK ----------
        st.markdown("### 📌 Fact Check")
        fact_placeholder = st.empty()
    
        def render_grammar(raw_text):
            if render_ai_error("Spelling/Grammar AI", raw_text, spelling_placeholder):
                render_ai_error("Spelling/Grammar AI", raw_text, grammar_placeholder)
                return
            clean = filter_invalid_rows(raw_text, article_text)
            spelling_table, grammar_table = split_spelling_grammar(clean)
            if spelling_table:
                spelling_placeholder.markdown(
                    render_language_table(spelling_table),
                    unsafe_allow_html=True,
                )
            else:
                spelling_placeholder.success("✅ No spelling issues found")
            if grammar_table:
                grammar_placeholder.markdown(
                    render_language_table(grammar_table),
                    unsafe_allow_html=True,
                )
            else:
                grammar_placeholder.success("✅ No grammar issues found")
    
        def render_fact(fact_text):
            if render_ai_error("Fact-check AI", fact_text, fact_placeholder):
                return
            if not fact_text or "| Statement |" not in fact_text:
                fact_placeholder.success("✅ No factual issues found")
            else:
                fact_placeholder.markdown(fact_text)
    
        # Show cached results if present
        if "grammar_raw" in results:
            render_grammar(results["grammar_raw"])
        else:
            spelling_placeholder.info("Running Gemini grammar/spelling...")
            grammar_placeholder.info("Running Gemini grammar/spelling...")
    
        if run_gemini_editorial:
            if "gemini_editorial" in results:
                if not render_ai_error("Editorial AI", results["gemini_editorial"], editorial_placeholder):
                    editorial_placeholder.markdown(results["gemini_editorial"])
            else:
                editorial_placeholder.info("Running Gemini editorial QC...")
    
        if "fact_result" in results:
            render_fact(results["fact_result"])
        else:
            fact_placeholder.info("Running Gemini fact check...")
    
        # Run missing AI tasks in parallel
        tasks = {}
        with ThreadPoolExecutor(max_workers=3) as executor:
            if "grammar_raw" not in results:
                tasks[executor.submit(
                    cached_gemini_grammar_review,
                    qc_content,
                    GRAMMAR_MAX_CHARS
                )] = "grammar"
            if run_gemini_editorial and "gemini_editorial" not in results:
                tasks[executor.submit(
                    cached_gemini_editorial_review,
                    qc_content,
                    web_story,
                    health_beauty
                )] = "editorial"
            if "fact_result" not in results:
                tasks[executor.submit(
                    cached_gemini_fact_check,
                    qc_content,
                    FACT_MAX_CHARS,
                    FACT_MAX_ITEMS
                )] = "fact"
    
            for future in as_completed(tasks):
                key = tasks[future]
                try:
                    result = future.result()
                except Exception as exc:
                    result = format_ai_error(key, exc)
    
                if key == "grammar":
                    results["grammar_raw"] = result
                    render_grammar(result)
                elif key == "editorial":
                    results["gemini_editorial"] = result
                    if editorial_placeholder:
                        if not render_ai_error("Editorial AI", result, editorial_placeholder):
                            editorial_placeholder.markdown(result)
                elif key == "fact":
                    results["fact_result"] = result
                    render_fact(result)
    
        clean_grammar_md = filter_invalid_rows(results.get("grammar_raw", ""), article_text) if "grammar_raw" in results else ""
        spelling_md, grammar_md = split_spelling_grammar(clean_grammar_md) if clean_grammar_md else ("", "")
        spelling_count = count_markdown_rows(spelling_md, "Original")
        grammar_count = count_markdown_rows(grammar_md, "Original")
        editorial_count = count_markdown_rows(results.get("gemini_editorial", ""), "Issue") if run_gemini_editorial else 0
        fact_count = count_markdown_rows(results.get("fact_result", ""), "Statement")
        has_ai_error = any(
            is_ai_error_output(value)
            for value in (
                results.get("grammar_raw", ""),
                results.get("gemini_editorial", "") if run_gemini_editorial else "",
                results.get("fact_result", ""),
            )
        )
        with score_placeholder.container():
            render_qc_score_summary(
                spelling_count,
                grammar_count,
                editorial_count,
                fact_count,
                has_ai_error,
                count_article_words(article_content),
            )
    
        report_pdf_bytes, report_pdf_error = build_english_qc_report_pdf(
            source_label or (url.strip() if source == "URL" and url else current_key or "QC Report"),
            _current_access_email(),
            spelling_md,
            grammar_md,
            results.get("gemini_editorial", "") if run_gemini_editorial else "",
            results.get("fact_result", ""),
        )
    
        if report_pdf_bytes:
            st.download_button(
                "Download QC Report (PDF)",
                data=report_pdf_bytes,
                file_name=f"english_qc_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf",
                mime="application/pdf",
            )
        elif report_pdf_error:
            st.caption(f"PDF report unavailable: {report_pdf_error}")
    
        required_keys = {"grammar_raw", "fact_result"}
        if run_gemini_editorial:
            required_keys.add("gemini_editorial")
    
        if "elapsed" not in results and required_keys.issubset(results.keys()):
            results["elapsed"] = time.perf_counter() - results["analysis_start"]
            log_analysis_run(
                "english_qc",
                _current_access_email(),
                st.session_state.get("_pending_source_type", source.lower()),
                st.session_state.get("_pending_source_identity", current_key or ""),
                st.session_state.get("_pending_source_label", url.strip() if source == "URL" and url else ""),
                st.session_state.get("_pending_analysis_key", analysis_key or ""),
                spelling_count,
                grammar_count,
                editorial_count,
                fact_count,
            )
    
        elapsed = results.get("elapsed")
        if elapsed is not None:
            minutes = int(elapsed // 60)
            seconds = int(round(elapsed % 60))
            st.caption(f"⏱️ Time taken: {minutes}m {seconds:02d}s")
            st.sidebar.metric("Time taken", f"{minutes}m {seconds:02d}s")
    
    if _is_admin_user():
        render_admin_dashboard("english_qc")
    
